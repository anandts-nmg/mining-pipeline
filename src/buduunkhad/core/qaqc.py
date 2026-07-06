"""QA/QC log model and writers.

Invariant #6: every phase writes a QA/QC log with item / acceptance / reviewer /
date / decision, and ends at a decision gate. The log is written as .xlsx (and
optionally .docx) so it slots into the methodology's handover package.
"""

from __future__ import annotations

from dataclasses import dataclass, field
from datetime import date
from enum import StrEnum
from pathlib import Path


class Decision(StrEnum):
    """Per-item QA/QC outcome."""

    PASS = "Pass"
    FAIL = "Fail"
    NA = "N/A"
    PENDING = "Pending"


@dataclass
class QAQCItem:
    """One QA/QC checklist row."""

    item: str
    acceptance: str
    decision: Decision = Decision.PENDING
    reviewer: str = ""
    reviewed_on: str = ""  # ISO date string; empty until reviewed
    note: str = ""

    def as_row(self) -> dict[str, str]:
        return {
            "item": self.item,
            "acceptance": self.acceptance,
            "decision": self.decision.value,
            "reviewer": self.reviewer,
            "date": self.reviewed_on,
            "note": self.note,
        }


@dataclass
class QAQCReport:
    """A phase's full QA/QC log."""

    phase_id: str
    phase_name: str
    items: list[QAQCItem] = field(default_factory=list)
    created_on: str = ""

    def add(
        self,
        item: str,
        acceptance: str,
        *,
        decision: Decision = Decision.PENDING,
        reviewer: str = "",
        note: str = "",
        reviewed_on: str | None = None,
    ) -> QAQCItem:
        rec = QAQCItem(
            item=item,
            acceptance=acceptance,
            decision=decision,
            reviewer=reviewer,
            reviewed_on=reviewed_on or "",
            note=note,
        )
        self.items.append(rec)
        return rec

    @property
    def has_failures(self) -> bool:
        return any(i.decision is Decision.FAIL for i in self.items)

    @property
    def has_pending(self) -> bool:
        return any(i.decision is Decision.PENDING for i in self.items)

    @property
    def passed(self) -> bool:
        """True if there are items and none failed. **PENDING is allowed** — an
        orchestrate phase legitimately ends with human-completion items still open.

        So ``passed`` means "nothing failed", *not* "all work complete". A consumer
        that needs true completion must also check :pyattr:`has_pending` (surfaced as
        ``qaqc_pending`` in the run manifest, and as the gate's ``provisional`` flag).
        """
        return bool(self.items) and not self.has_failures

    # ---- writers --------------------------------------------------------- #

    def write_xlsx(self, path: Path) -> Path:
        from openpyxl import Workbook
        from openpyxl.styles import Font

        path = Path(path)
        path.parent.mkdir(parents=True, exist_ok=True)
        wb = Workbook()
        ws = wb.active
        assert ws is not None  # a fresh Workbook always has an active worksheet
        ws.title = f"Phase {self.phase_id} QAQC"
        header = ["QA/QC item", "Acceptance", "Decision", "Reviewer", "Date", "Note"]
        ws.append(header)
        for cell in ws[1]:
            cell.font = Font(bold=True)
        for it in self.items:
            r = it.as_row()
            ws.append(
                [r["item"], r["acceptance"], r["decision"], r["reviewer"], r["date"], r["note"]]
            )
        widths = [42, 48, 12, 16, 12, 40]
        for idx, w in enumerate(widths, start=1):
            ws.column_dimensions[chr(64 + idx)].width = w
        wb.save(path)
        return path

    def write_docx(self, path: Path) -> Path:
        from docx import Document

        path = Path(path)
        path.parent.mkdir(parents=True, exist_ok=True)
        doc = Document()
        doc.add_heading(f"Phase {self.phase_id} — {self.phase_name}", level=1)
        doc.add_heading("QA/QC Log", level=2)
        if self.created_on:
            doc.add_paragraph(f"Created: {self.created_on}")
        table = doc.add_table(rows=1, cols=6)
        table.style = "Light Grid Accent 1"
        hdr = table.rows[0].cells
        for i, label in enumerate(
            ["QA/QC item", "Acceptance", "Decision", "Reviewer", "Date", "Note"]
        ):
            hdr[i].text = label
        for it in self.items:
            r = it.as_row()
            cells = table.add_row().cells
            cells[0].text = r["item"]
            cells[1].text = r["acceptance"]
            cells[2].text = r["decision"]
            cells[3].text = r["reviewer"]
            cells[4].text = r["date"]
            cells[5].text = r["note"]
        doc.save(str(path))
        return path


# Standard acceptance wording from the methodology.
RECORDED_ACCEPTANCE = "Recorded in phase QA/QC log; reviewer/date/decision required."


def new_report(phase_id: str, phase_name: str, *, created_on: str | None = None) -> QAQCReport:
    return QAQCReport(
        phase_id=phase_id,
        phase_name=phase_name,
        created_on=created_on or date.today().isoformat(),
    )
