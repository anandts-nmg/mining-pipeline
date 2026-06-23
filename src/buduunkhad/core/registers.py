"""Register / inventory / readme writers (xlsx, csv, docx).

These produce the methodology's bookkeeping artefacts: the 78-input master
inventory, the raw-data integrity log, the SHA-256 checksum register, the data
confidence ranking, and the source-data readme.
"""

from __future__ import annotations

import csv
from collections.abc import Iterable, Sequence
from pathlib import Path

from buduunkhad.core.raw_guard import ChecksumRecord

# Master inventory columns (methodology Phase 00 step 2 + section 4.1).
INVENTORY_COLUMNS: list[str] = [
    "no",
    "evidence_group",
    "primary_phase",
    "original_filename",
    "standardized_filename",
    "file_type",
    "is_sidecar",
    "parent_file",
    "raw_path",
    "working_copy_path",
    "file_size_mb",
    "checksum_sha256",
    "read_status",
    "open_status",
    "processing_copy_status",
    "scan_quality",
    "has_coordinate_grid",
    "georef_required",
    "georef_status",
    "vectorization_status",
    "handover_status",
    "source_note",
    "owner",
    "methodology_action",
]

CONFIDENCE_COLUMNS: list[str] = [
    "no",
    "evidence_group",
    "filename",
    "primary_phase",
    "data_confidence",  # High / Medium / Low / Needs verification
    "evidence_role",  # support / decision
    "limitation_note",
    "reviewer",
    "date",
]


def _autosize(ws, widths: Sequence[int]) -> None:
    for idx, w in enumerate(widths, start=1):
        col = ws.cell(row=1, column=idx).column_letter
        ws.column_dimensions[col].width = w


def write_table_xlsx(
    rows: Iterable[dict[str, object]],
    columns: list[str],
    path: Path,
    *,
    sheet_title: str = "Sheet1",
    widths: Sequence[int] | None = None,
) -> Path:
    """Generic typed-header xlsx writer used by the register helpers."""
    from openpyxl import Workbook
    from openpyxl.styles import Font

    path = Path(path)
    path.parent.mkdir(parents=True, exist_ok=True)
    wb = Workbook()
    ws = wb.active
    assert ws is not None  # a fresh Workbook always has an active worksheet
    ws.title = sheet_title[:31]
    ws.append(columns)
    for cell in ws[1]:
        cell.font = Font(bold=True)
    for row in rows:
        ws.append([row.get(c, "") for c in columns])
    _autosize(ws, widths or [18] * len(columns))
    ws.freeze_panes = "A2"
    wb.save(path)
    return path


def write_inventory_xlsx(rows: Iterable[dict[str, object]], path: Path) -> Path:
    """Write the 78-input master inventory."""
    return write_table_xlsx(
        rows,
        INVENTORY_COLUMNS,
        path,
        sheet_title="Master Inventory",
        widths=[
            5,
            28,
            8,
            44,
            44,
            12,
            9,
            30,
            40,
            40,
            11,
            24,
            11,
            14,
            18,
            12,
            16,
            14,
            16,
            18,
            16,
            30,
            14,
            44,
        ],
    )


def write_confidence_ranking_xlsx(rows: Iterable[dict[str, object]], path: Path) -> Path:
    """Write the data confidence ranking."""
    return write_table_xlsx(
        rows,
        CONFIDENCE_COLUMNS,
        path,
        sheet_title="Data Confidence",
        widths=[5, 28, 44, 8, 18, 12, 40, 16, 12],
    )


def write_integrity_log_xlsx(records: Sequence[ChecksumRecord], path: Path) -> Path:
    """Write the raw-data integrity log (checksums + size, one row per raw file)."""
    rows = [
        {
            "filename": r.filename,
            "relative_path": r.relative_path,
            "sha256": r.sha256,
            "size_bytes": r.size_bytes,
            "size_mb": round(r.size_bytes / (1024 * 1024), 4),
        }
        for r in records
    ]
    return write_table_xlsx(
        rows,
        ["filename", "relative_path", "sha256", "size_bytes", "size_mb"],
        path,
        sheet_title="Raw Integrity Log",
        widths=[44, 50, 66, 14, 12],
    )


def write_checksum_register_csv(records: Sequence[ChecksumRecord], path: Path) -> Path:
    """Write the SHA-256 checksum register as CSV."""
    path = Path(path)
    path.parent.mkdir(parents=True, exist_ok=True)
    with path.open("w", encoding="utf-8", newline="") as f:
        w = csv.writer(f)
        w.writerow(["filename", "relative_path", "sha256", "size_bytes"])
        for r in records:
            w.writerow([r.filename, r.relative_path, r.sha256, r.size_bytes])
    return path


def read_checksum_register_csv(path: Path) -> dict[str, str]:
    """Read a checksum register back into ``{relative_path: sha256}``."""
    path = Path(path)
    out: dict[str, str] = {}
    if not path.exists():
        return out
    with path.open(encoding="utf-8", newline="") as f:
        for row in csv.DictReader(f):
            out[row["relative_path"]] = row["sha256"]
    return out


def write_source_readme_docx(
    path: Path,
    *,
    title: str,
    project: str,
    license_code: str,
    target_crs: str,
    evidence_groups: Sequence[tuple[str, int]],
    notes: Sequence[str] = (),
) -> Path:
    """Write the source-data readme describing the raw archive."""
    from docx import Document

    path = Path(path)
    path.parent.mkdir(parents=True, exist_ok=True)
    doc = Document()
    doc.add_heading(title, level=1)
    doc.add_paragraph(f"Project: {project}")
    doc.add_paragraph(f"License: {license_code}")
    doc.add_paragraph(f"Standard deliverable CRS: {target_crs}")

    doc.add_heading("Evidence groups (78 raw inputs)", level=2)
    table = doc.add_table(rows=1, cols=2)
    table.style = "Light Grid Accent 1"
    table.rows[0].cells[0].text = "Evidence group"
    table.rows[0].cells[1].text = "File count"
    for name, count in evidence_groups:
        cells = table.add_row().cells
        cells[0].text = name
        cells[1].text = str(count)

    doc.add_heading("Governing principles", level=2)
    default_notes = [
        "Raw data is read-only: never modified, renamed, moved, clipped or reprojected.",
        "All processing happens on working copies under the output root.",
        "Sidecar metadata (.tfw .jgw .aux.xml .ovr .rpc .eph .txt) travels with its parent.",
        "All deliverables are EPSG:32647; native/source CRS is recorded, never dropped.",
        "Remote sensing / pXRF / drone are support evidence; lab assay + field geology + "
        "structural control are decision evidence.",
    ]
    for note in [*default_notes, *notes]:
        doc.add_paragraph(note, style="List Bullet")

    doc.save(str(path))
    return path
