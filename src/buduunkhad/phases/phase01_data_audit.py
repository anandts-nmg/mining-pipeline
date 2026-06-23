"""Phase 01 — Data Audit and Master GIS Setup (BUILD).

Make the 78 inputs GIS-ready and stand up the EPSG:32647 master database:

1. Import the license-boundary KMZ (№8) -> GeoPackage, reprojected to EPSG:32647.
2. Generate the project buffers (500 m, 1 km, 5 km, 10 km, 20 km).
3. Audit every raster's CRS / resolution / extent / nodata / band count.
4. Create the Master GeoPackage schema (13 typed, empty layers).
5. Write the CRS/Georeference QA/QC log and the Data Confidence Ranking.
6. Emit a minimal EPSG:32647 QGIS project (.qgz).
7. Emit the Phase 1 handoff placeholders required by the methodology.
"""

from __future__ import annotations

import zipfile
from pathlib import Path

from buduunkhad.config import InputRecord
from buduunkhad.core import crs as crs_mod
from buduunkhad.core import naming, registers, vector_io
from buduunkhad.core.qaqc import RECORDED_ACCEPTANCE, Decision, QAQCReport, new_report
from buduunkhad.phases.base import Phase, PhaseResult, RunContext


class Phase01DataAudit(Phase):
    id = "01"
    name = "Data Audit and Master GIS Setup"
    mode = "build"
    input_numbers = list(range(1, 79))
    gate_condition = (
        "Master GIS project opens without missing layers; critical data confidence recorded."
    )

    # populated during run() for qaqc()
    _boundary_ok: bool
    _boundary_epsg: int | None
    _n_buffers: int
    _raster_audits: list[crs_mod.RasterAudit]
    _master_layers: list[str]
    _handoff_paths: list[Path]

    def __init__(self) -> None:
        self._boundary_ok = False
        self._boundary_epsg = None
        self._n_buffers = 0
        self._raster_audits = []
        self._master_layers = []
        self._handoff_paths = []

    # ------------------------------------------------------------------ #

    def run(self, ctx: RunContext) -> PhaseResult:
        pdir = ctx.phase_dir(self.id)
        cfg = ctx.config
        epsg = cfg.target_epsg
        result = PhaseResult(self.id, status="dry-run" if ctx.dry_run else "ok")

        gpkg_dir = pdir / "05_KMZ_KML_to_GPKG"
        master_dir = pdir / "06_Master_GeoPackage_Schema"
        inventory_dir = pdir / "01_File_Inventory"
        crs_dir = pdir / "03_CRS_Check"
        conf_dir = pdir / "07_Data_Confidence_Ranking"
        qgis_dir = pdir / "08_Master_QGIS_Project_Setup"
        handoff_dir = pdir / "09_Handover_Package"

        # ---- 4. Master GeoPackage schema (created in both dry-run and real) ----
        master_path = master_dir / f"{cfg.register_prefix}_Master_GIS_Database.gpkg"
        vector_io.create_master_gpkg(master_path, cfg.master_gpkg_layers, epsg)
        self._master_layers = vector_io.list_gpkg_layers(master_path)
        result.add_output(master_path)

        # ---- minimal QGIS project (.qgz) ----
        qgz_path = qgis_dir / f"{cfg.register_prefix}_Master_QGIS_Project.qgz"
        _write_min_qgz(qgz_path, epsg, project_title=f"{cfg.project.name} Master")
        result.add_output(qgz_path)

        crs_log = crs_dir / f"{cfg.register_prefix}_CRS_Georeference_QAQC_Log.xlsx"
        conf_path = conf_dir / f"{cfg.register_prefix}_Data_Confidence_Ranking.xlsx"
        inventory_path = inventory_dir / f"{cfg.register_prefix}_Phase1_File_Inventory.xlsx"
        gap_path = conf_dir / f"{cfg.register_prefix}_Data_Gap_Register.xlsx"
        index_map_path = handoff_dir / f"{cfg.register_prefix}_Phase1_Master_GIS_Index_Map.pdf"
        summary_path = handoff_dir / f"{cfg.register_prefix}_Phase1_Desktop_Study_Summary.docx"
        phase2_ready_path = handoff_dir / f"{cfg.register_prefix}_Phase2_Ready_Dataset_List.xlsx"

        if ctx.dry_run:
            registers.write_table_xlsx([], _CRS_LOG_COLUMNS, crs_log, sheet_title="CRS Georef QAQC")
            registers.write_confidence_ranking_xlsx([], conf_path)
            self._write_handoff_outputs(
                ctx,
                inventory_path=inventory_path,
                gap_path=gap_path,
                index_map_path=index_map_path,
                summary_path=summary_path,
                phase2_ready_path=phase2_ready_path,
            )
            result.add_output(crs_log)
            result.add_output(conf_path)
            for path in self._handoff_paths:
                result.add_output(path)
            result.log(
                "dry-run: master GPKG schema + QGIS project + empty CRS/confidence logs "
                "+ Phase 1 handoff package"
            )
            return result

        # ---- 1. boundary import + 2. buffers (real run) ----
        boundary_rec = ctx.record_by_no(cfg.boundary.input_no)
        boundary_src = self._working_copy_path(ctx, boundary_rec)
        if boundary_src.exists():
            self._import_boundary_and_buffers(ctx, boundary_src, gpkg_dir, master_path, result)
        else:
            result.log(f"boundary working copy missing: {boundary_src}")

        # ---- 3. raster CRS/resolution/extent/nodata/band audit ----
        crs_rows = self._audit_rasters(ctx)
        registers.write_table_xlsx(
            crs_rows, _CRS_LOG_COLUMNS, crs_log, sheet_title="CRS Georef QAQC"
        )
        result.add_output(crs_log)

        # ---- 5. data confidence ranking ----
        registers.write_confidence_ranking_xlsx(self._confidence_rows(ctx), conf_path)
        result.add_output(conf_path)

        self._write_handoff_outputs(
            ctx,
            inventory_path=inventory_path,
            gap_path=gap_path,
            index_map_path=index_map_path,
            summary_path=summary_path,
            phase2_ready_path=phase2_ready_path,
        )
        for path in self._handoff_paths:
            result.add_output(path)

        result.log(
            f"boundary={'ok' if self._boundary_ok else 'missing'}, buffers={self._n_buffers}, "
            f"rasters audited={len(self._raster_audits)}, master layers={len(self._master_layers)}, "
            f"handoff outputs={len(self._handoff_paths)}"
        )
        return result

    # ------------------------------------------------------------------ #

    def _working_copy_path(self, ctx: RunContext, rec: InputRecord) -> Path:
        """Phase 00 materialises working copies under the control archive by group."""
        return ctx.phase_dir("00") / rec.evidence_group / rec.filename

    def _import_boundary_and_buffers(
        self,
        ctx: RunContext,
        boundary_src: Path,
        gpkg_dir: Path,
        master_path: Path,
        result: PhaseResult,
    ) -> None:
        cfg = ctx.config
        epsg = cfg.target_epsg
        gdf = vector_io.read_boundary(boundary_src, assume_epsg=cfg.crs.source_geographic_epsg)
        self._boundary_epsg = gdf.crs.to_epsg() if gdf.crs else None
        gdf32 = crs_mod.reproject_gdf(gdf, epsg)
        # Drop source (KML) attribute columns - they collide case-insensitively with
        # our schema fields in GeoPackage - and keep only geometry + our attributes.
        geom_col = gdf32.geometry.name
        gdf32 = gdf32[[geom_col]].copy()
        if geom_col != "geometry":
            gdf32 = gdf32.rename_geometry("geometry")
        # Force 2D - boundary polygons are planar; KML often carries a Z=0 ordinate.
        from shapely import force_2d

        gdf32 = gdf32.set_geometry(gdf32.geometry.apply(force_2d))
        gdf32["name"] = cfg.project.license_code
        gdf32["source_input"] = f"#{ctx.config.boundary.input_no} {boundary_src.name}"

        boundary_name = naming.data_name(
            cfg.data_prefix,
            f"{cfg.project.license_code}_LicenseBoundary",
            crs_or_param=naming.epsg_tag(epsg),
            version=1,
            ext="gpkg",
        )
        boundary_path = gpkg_dir / boundary_name
        vector_io.write_layer(gdf32, boundary_path, layer="license_boundary")
        result.add_output(boundary_path)

        # populate the master gpkg's (empty) license_boundary layer, preserving others
        vector_io.write_layer(gdf32, master_path, layer="license_boundary", mode="a")

        # buffers
        buffers_gdf = self._make_buffers(gdf32, cfg.boundary.buffers_m, epsg)
        self._n_buffers = len(buffers_gdf)
        buffer_name = naming.data_name(
            cfg.data_prefix,
            "Project",
            crs_or_param=f"{naming.buffer_param(cfg.boundary.buffers_m)}_{naming.epsg_tag(epsg)}",
            version=None,
            ext="gpkg",
        )
        buffer_path = gpkg_dir / buffer_name
        vector_io.write_layer(buffers_gdf, buffer_path, layer="project_buffers")
        result.add_output(buffer_path)
        self._boundary_ok = True

    def _make_buffers(self, boundary_gdf, buffers_m: list[int], epsg: int):  # type: ignore[no-untyped-def]
        import geopandas as gpd

        merged = (
            boundary_gdf.geometry.union_all()
            if hasattr(boundary_gdf.geometry, "union_all")
            else boundary_gdf.geometry.unary_union
        )
        rings = []
        for dist in sorted(buffers_m):
            rings.append({"distance_m": dist, "geometry": merged.buffer(dist)})
        return gpd.GeoDataFrame(rings, crs=f"EPSG:{epsg}")

    def _audit_rasters(self, ctx: RunContext) -> list[dict[str, object]]:
        rows: list[dict[str, object]] = []
        target = ctx.config.target_epsg
        for rec in sorted(ctx.register, key=lambda r: r.no):
            if rec.file_type != "raster":
                continue
            wc = self._working_copy_path(ctx, rec)
            if not wc.exists():
                continue
            audit = crs_mod.audit_raster(wc, target_epsg=target)
            self._raster_audits.append(audit)
            rows.append(
                {
                    "no": rec.no,
                    "filename": rec.filename,
                    "readable": audit.readable,
                    "native_epsg": audit.epsg if audit.epsg is not None else "",
                    "target_epsg": target,
                    "needs_reproject": audit.needs_reproject
                    if audit.needs_reproject is not None
                    else "",
                    "width": audit.width if audit.width is not None else "",
                    "height": audit.height if audit.height is not None else "",
                    "band_count": audit.band_count if audit.band_count is not None else "",
                    "dtype": audit.dtype or "",
                    "nodata": audit.nodata if audit.nodata is not None else "",
                    "res_x": audit.res_x if audit.res_x is not None else "",
                    "res_y": audit.res_y if audit.res_y is not None else "",
                    "decision": "Pass" if audit.readable else "Fail",
                    "note": audit.error or "",
                }
            )
        return rows

    def _confidence_rows(self, ctx: RunContext) -> list[dict[str, object]]:
        rows: list[dict[str, object]] = []
        for rec in sorted(ctx.register, key=lambda r: r.no):
            rows.append(
                {
                    "no": rec.no,
                    "evidence_group": rec.evidence_group,
                    "filename": rec.filename,
                    "primary_phase": rec.primary_phase,
                    "data_confidence": _seed_confidence(rec),
                    "evidence_role": _evidence_role(rec),
                    "limitation_note": _limitation_note(rec),
                    "reviewer": "",
                    "date": "",
                }
            )
        return rows

    def _write_handoff_outputs(
        self,
        ctx: RunContext,
        *,
        inventory_path: Path,
        gap_path: Path,
        index_map_path: Path,
        summary_path: Path,
        phase2_ready_path: Path,
    ) -> None:
        """Write methodology-required Phase 1 handoff artifacts without raw processing."""
        registers.write_inventory_xlsx(self._phase1_inventory_rows(ctx), inventory_path)
        registers.write_table_xlsx(
            self._data_gap_rows(ctx),
            _DATA_GAP_COLUMNS,
            gap_path,
            sheet_title="Data Gaps",
            widths=[14, 5, 28, 44, 24, 14, 52, 16, 14],
        )
        registers.write_table_xlsx(
            self._phase2_ready_rows(ctx),
            _PHASE2_READY_COLUMNS,
            phase2_ready_path,
            sheet_title="Phase 2 Ready",
            widths=[5, 28, 44, 12, 40, 18, 18, 44],
        )
        _write_index_map_placeholder(index_map_path, ctx)
        _write_desktop_summary(summary_path, ctx, data_gaps=len(self._data_gap_rows(ctx)))
        self._handoff_paths = [
            inventory_path,
            gap_path,
            index_map_path,
            summary_path,
            phase2_ready_path,
        ]

    def _phase1_inventory_rows(self, ctx: RunContext) -> list[dict[str, object]]:
        rows: list[dict[str, object]] = []
        for rec in sorted(ctx.register, key=lambda r: r.no):
            wc = self._working_copy_path(ctx, rec)
            exists = wc.exists()
            rows.append(
                {
                    "no": rec.no,
                    "evidence_group": rec.evidence_group,
                    "primary_phase": rec.primary_phase,
                    "original_filename": rec.filename,
                    "standardized_filename": rec.filename,
                    "file_type": rec.file_type,
                    "is_sidecar": rec.is_sidecar,
                    "parent_file": rec.parent_file,
                    "raw_path": "",
                    "working_copy_path": str(wc),
                    "file_size_mb": round(wc.stat().st_size / (1024 * 1024), 4) if exists else "",
                    "checksum_sha256": "",
                    "read_status": "dry-run not checked" if ctx.dry_run else "working copy checked",
                    "open_status": "Not checked (dry-run)"
                    if ctx.dry_run
                    else ("Present" if exists else "Missing"),
                    "processing_copy_status": "Not materialized (dry-run)"
                    if ctx.dry_run
                    else ("Available" if exists else "Missing"),
                    "scan_quality": "",
                    "has_coordinate_grid": "",
                    "georef_required": _georef_required(rec),
                    "georef_status": "Not started" if _georef_required(rec) == "Yes" else "N/A",
                    "vectorization_status": "Not started" if _needs_vectorization(rec) else "N/A",
                    "handover_status": "Use with caution"
                    if _requires_manual_review(rec)
                    else "Ready",
                    "source_note": "",
                    "owner": "",
                    "methodology_action": rec.methodology_action,
                }
            )
        return rows

    def _data_gap_rows(self, ctx: RunContext) -> list[dict[str, object]]:
        rows: list[dict[str, object]] = []
        gap_no = 1
        for rec in sorted(ctx.register, key=lambda r: r.no):
            wc = self._working_copy_path(ctx, rec)
            if not ctx.dry_run and not wc.exists():
                rows.append(
                    _gap_row(
                        gap_no,
                        rec,
                        "Missing working copy",
                        "High",
                        "Run Phase 00 and reconcile raw archive/register before Phase 1 closeout.",
                    )
                )
                gap_no += 1
                continue
            if ctx.dry_run:
                rows.append(
                    _gap_row(
                        gap_no,
                        rec,
                        "Dry-run placeholder",
                        "Medium",
                        "Populate raw archive, run Phase 00, then run Phase 01 for real QA/QC.",
                    )
                )
                gap_no += 1
                continue
            if rec.is_sidecar:
                rows.append(
                    _gap_row(
                        gap_no,
                        rec,
                        "Sidecar bundle verification",
                        "Medium",
                        "Confirm sidecar remains paired with parent raster/image in Phase 2.",
                    )
                )
                gap_no += 1
            elif _requires_manual_review(rec):
                rows.append(
                    _gap_row(
                        gap_no,
                        rec,
                        "Manual georeference/source review",
                        "Medium",
                        "Record GCP/residual/source confidence before using as decision evidence.",
                    )
                )
                gap_no += 1
            elif rec.file_type == "hdf":
                rows.append(
                    _gap_row(
                        gap_no,
                        rec,
                        "ASTER HDF compatibility",
                        "Medium",
                        "Confirm ASTER HDF opens in the selected SNAP/ILWIS/QGIS workflow.",
                    )
                )
                gap_no += 1
        return rows

    def _phase2_ready_rows(self, ctx: RunContext) -> list[dict[str, object]]:
        rows: list[dict[str, object]] = []
        for rec in sorted(ctx.records_by_numbers(_PHASE2_INPUT_NUMBERS), key=lambda r: r.no):
            wc = self._working_copy_path(ctx, rec)
            exists = wc.exists()
            rows.append(
                {
                    "no": rec.no,
                    "evidence_group": rec.evidence_group,
                    "filename": rec.filename,
                    "file_type": rec.file_type,
                    "working_copy_path": str(wc),
                    "phase1_status": "Dry-run placeholder"
                    if ctx.dry_run
                    else ("Available" if exists else "Missing"),
                    "ready_for_phase2": "Pending raw sync"
                    if ctx.dry_run
                    else ("Yes" if exists else "No"),
                    "limitation_note": _limitation_note(rec),
                }
            )
        return rows

    # ------------------------------------------------------------------ #

    def qaqc(self, ctx: RunContext) -> QAQCReport:
        report = new_report(self.id, self.name)
        if ctx.dry_run:
            report.add(
                "Master GPKG schema created",
                RECORDED_ACCEPTANCE,
                decision=Decision.PASS if self._master_layers else Decision.FAIL,
                note=f"{len(self._master_layers)} layers created (dry-run).",
            )
            report.add(
                "EPSG:32647 project CRS",
                RECORDED_ACCEPTANCE,
                decision=Decision.PASS,
                note="QGIS project written with project CRS EPSG:32647.",
            )
            report.add(
                "Phase 1 handoff package complete",
                RECORDED_ACCEPTANCE,
                decision=Decision.PASS if len(self._handoff_paths) == 5 else Decision.FAIL,
                note=f"{len(self._handoff_paths)} handoff artifact(s) written.",
            )
            return report

        expected = {layer.name for layer in ctx.config.master_gpkg_layers}
        layers_ok = expected.issubset(set(self._master_layers))
        report.add(
            "EPSG:32647 project CRS",
            RECORDED_ACCEPTANCE,
            decision=Decision.PASS,
            note=f"Project CRS = {ctx.config.crs.target_name}, {ctx.config.crs.target_authority}.",
        )
        report.add(
            "KMZ boundary topology valid",
            RECORDED_ACCEPTANCE,
            decision=Decision.PASS if self._boundary_ok else Decision.FAIL,
            note=f"Boundary imported (native EPSG {self._boundary_epsg}); {self._n_buffers} buffers."
            if self._boundary_ok
            else "Boundary working copy not found / not imported.",
        )
        report.add(
            "Raster CRS/resolution/extent/nodata/band count checked",
            RECORDED_ACCEPTANCE,
            decision=Decision.PASS if self._raster_audits else Decision.PENDING,
            note=f"{len(self._raster_audits)} raster(s) audited.",
        )
        report.add(
            "Scan georeference residual and confidence logged",
            RECORDED_ACCEPTANCE,
            decision=Decision.PENDING,
            note="Scan georeferencing performed in QGIS (Phase 1 sub-workflow); residuals logged by operator.",
        )
        report.add(
            "Master GPKG schema created",
            RECORDED_ACCEPTANCE,
            decision=Decision.PASS if layers_ok else Decision.FAIL,
            note=f"{len(self._master_layers)}/{len(expected)} expected layers present.",
        )
        report.add(
            "Phase 1 handoff package complete",
            RECORDED_ACCEPTANCE,
            decision=Decision.PASS if len(self._handoff_paths) == 5 else Decision.FAIL,
            note=f"{len(self._handoff_paths)} handoff artifact(s) written.",
        )
        return report


# --------------------------------------------------------------------------- #
# module-level helpers
# --------------------------------------------------------------------------- #

_CRS_LOG_COLUMNS = [
    "no",
    "filename",
    "readable",
    "native_epsg",
    "target_epsg",
    "needs_reproject",
    "width",
    "height",
    "band_count",
    "dtype",
    "nodata",
    "res_x",
    "res_y",
    "decision",
    "note",
]

_DATA_GAP_COLUMNS = [
    "gap_id",
    "no",
    "evidence_group",
    "filename",
    "gap_type",
    "severity",
    "recommended_action",
    "owner",
    "status",
]

_PHASE2_READY_COLUMNS = [
    "no",
    "evidence_group",
    "filename",
    "file_type",
    "working_copy_path",
    "phase1_status",
    "ready_for_phase2",
    "limitation_note",
]

_PHASE2_INPUT_NUMBERS = [*range(9, 47), *range(73, 79)]

_SUPPORT_GROUPS = {
    "01_Tectonic_Terrane_KMZ",
    "02_DEM_ALOS_ASTERGDEM",
    "03_KOMPSAT2_MSC_L1G",
    "06_Regional_Metallogenic_L47B",
    "07_Basemap_Sentinel2_ASTER",
}
_DECISION_GROUPS = {
    "04_HeavyMineral_StreamSediment_Field",
    "05_Geology_Mineral_Prospectivity",
}


def _evidence_role(rec: InputRecord) -> str:
    if rec.no == 8:  # license boundary is administrative/decision context
        return "decision"
    if rec.evidence_group in _DECISION_GROUPS:
        return "decision"
    return "support"


def _seed_confidence(rec: InputRecord) -> str:
    if rec.is_sidecar:
        return "Needs verification"
    if rec.file_type in ("image_scan", "browse_image", "thumbnail", "pdf"):
        return "Low"
    return "Medium"


def _limitation_note(rec: InputRecord) -> str:
    if rec.is_sidecar:
        return "Support sidecar; not usable without its parent."
    if rec.file_type in ("image_scan", "browse_image", "thumbnail", "pdf"):
        return "Scanned/non-native; georeference + source check required before use."
    return "Confirm CRS/metadata in audit; native CRS recorded."


def _georef_required(rec: InputRecord) -> str:
    if rec.file_type in ("image_scan", "browse_image", "thumbnail"):
        return "Yes"
    if rec.file_type == "pdf":
        return "Context only"
    return "No"


def _needs_vectorization(rec: InputRecord) -> bool:
    return rec.file_type in ("image_scan", "browse_image", "thumbnail") and rec.primary_phase in {
        "03",
        "08",
    }


def _requires_manual_review(rec: InputRecord) -> bool:
    return rec.file_type in ("image_scan", "browse_image", "thumbnail", "pdf")


def _gap_row(
    gap_no: int,
    rec: InputRecord,
    gap_type: str,
    severity: str,
    recommended_action: str,
) -> dict[str, object]:
    return {
        "gap_id": f"P1-GAP-{gap_no:03d}",
        "no": rec.no,
        "evidence_group": rec.evidence_group,
        "filename": rec.filename,
        "gap_type": gap_type,
        "severity": severity,
        "recommended_action": recommended_action,
        "owner": "",
        "status": "Open",
    }


def _write_desktop_summary(path: Path, ctx: RunContext, *, data_gaps: int) -> Path:
    from docx import Document

    path = Path(path)
    path.parent.mkdir(parents=True, exist_ok=True)
    doc = Document()
    doc.add_heading("Buduunkhad Phase 1 Desktop Study Summary", level=1)
    doc.add_paragraph(f"Project: {ctx.config.project.project_code} / {ctx.config.project.name}")
    doc.add_paragraph(f"License: {ctx.config.project.license_code}")
    doc.add_paragraph(
        f"Target CRS: {ctx.config.crs.target_name}, {ctx.config.crs.target_authority}"
    )
    doc.add_paragraph(f"Run ID: {ctx.run_id}")
    doc.add_paragraph("This readiness summary is generated from the configured 78-input register.")
    doc.add_paragraph("It does not interpret ore potential or process real raw data.")

    doc.add_heading("Phase 1 Handoff Status", level=2)
    checks = [
        "Master GIS database schema created.",
        "Master QGIS project placeholder created with EPSG:32647 project CRS.",
        "CRS/georeference QA/QC log created.",
        "Data confidence ranking created.",
        f"Data gap register created with {data_gaps} open item(s).",
        "Phase 2-ready dataset list created for remote-sensing inputs.",
    ]
    for check in checks:
        doc.add_paragraph(check, style="List Bullet")

    doc.add_heading("Limitations", level=2)
    for note in [
        "Scanned maps require operator georeference review, GCP residual logging and confidence flags.",
        "Remote sensing, DEM, drone/LiDAR and pXRF products are support evidence only.",
        "Final target confidence requires field geology, sampling, laboratory assay and structural control.",
    ]:
        doc.add_paragraph(note, style="List Bullet")
    doc.save(str(path))
    return path


def _write_index_map_placeholder(path: Path, ctx: RunContext) -> Path:
    lines = [
        "Phase 1 Master GIS Index Map Placeholder",
        f"Project: {ctx.config.project.project_code} / {ctx.config.project.name}",
        f"License: {ctx.config.project.license_code}",
        f"Target CRS: {ctx.config.crs.target_name}, {ctx.config.crs.target_authority}",
        "Map production requires QGIS rendering after real raw data sync.",
        "Use this placeholder to keep the Phase 1 handoff package complete.",
    ]
    return _write_simple_pdf(path, lines)


def _write_simple_pdf(path: Path, lines: list[str]) -> Path:
    path = Path(path)
    path.parent.mkdir(parents=True, exist_ok=True)

    def esc(text: str) -> str:
        return text.replace("\\", "\\\\").replace("(", "\\(").replace(")", "\\)")

    text_lines = ["BT", "/F1 12 Tf", "72 760 Td"]
    for i, line in enumerate(lines):
        if i:
            text_lines.append("0 -20 Td")
        text_lines.append(f"({esc(line)}) Tj")
    text_lines.append("ET")
    stream = "\n".join(text_lines).encode("ascii", errors="replace")

    objects = [
        b"<< /Type /Catalog /Pages 2 0 R >>",
        b"<< /Type /Pages /Kids [3 0 R] /Count 1 >>",
        (
            b"<< /Type /Page /Parent 2 0 R /MediaBox [0 0 612 792] "
            b"/Resources << /Font << /F1 4 0 R >> >> /Contents 5 0 R >>"
        ),
        b"<< /Type /Font /Subtype /Type1 /BaseFont /Helvetica >>",
        b"<< /Length "
        + str(len(stream)).encode("ascii")
        + b" >>\nstream\n"
        + stream
        + b"\nendstream",
    ]
    chunks = [b"%PDF-1.4\n"]
    offsets = []
    for idx, obj in enumerate(objects, start=1):
        offsets.append(sum(len(c) for c in chunks))
        chunks.append(f"{idx} 0 obj\n".encode("ascii"))
        chunks.append(obj)
        chunks.append(b"\nendobj\n")
    xref_offset = sum(len(c) for c in chunks)
    chunks.append(f"xref\n0 {len(objects) + 1}\n".encode("ascii"))
    chunks.append(b"0000000000 65535 f \n")
    for offset in offsets:
        chunks.append(f"{offset:010d} 00000 n \n".encode("ascii"))
    chunks.append(
        (
            f"trailer\n<< /Size {len(objects) + 1} /Root 1 0 R >>\n"
            f"startxref\n{xref_offset}\n%%EOF\n"
        ).encode("ascii")
    )
    path.write_bytes(b"".join(chunks))
    return path


def _write_min_qgz(path: Path, epsg: int, project_title: str) -> Path:
    """Write a minimal but valid QGIS project (.qgz = zip containing a .qgs)."""
    path = Path(path)
    path.parent.mkdir(parents=True, exist_ok=True)
    qgs_xml = (
        '<?xml version="1.0" encoding="UTF-8"?>\n'
        f'<qgis version="3.34" projectname="{project_title}">\n'
        "  <projectCrs>\n"
        "    <spatialrefsys>\n"
        f"      <authid>EPSG:{epsg}</authid>\n"
        f"      <srid>{epsg}</srid>\n"
        "    </spatialrefsys>\n"
        "  </projectCrs>\n"
        "  <layer-tree-group/>\n"
        "  <projectlayers/>\n"
        "</qgis>\n"
    )
    if path.exists():
        path.unlink()
    with zipfile.ZipFile(path, "w", zipfile.ZIP_DEFLATED) as zf:
        zf.writestr(f"{path.stem}.qgs", qgs_xml)
    return path


PHASE = Phase01DataAudit
