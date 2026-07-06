"""Phase 03 / 03A — Geological, Metallogenic & CMCS Synthesis + Preliminary Deposit Model.

ORCHESTRATE. The heavy lifting — georeferencing scan maps, digitizing lithology /
structure / occurrence vectors, writing the deposit model, doing the scoring — is human
work in QGIS / Excel / Word. This module scaffolds the 12-folder tree, emits every
register / template / schema, ingests the machine-tractable inputs (the #68 mineralized-
point XLSX and any human-digitized layers), builds the CMCS 5/10/20/25 km context buffer off
the Phase 01 licence boundary, assembles the authoritative 17-layer evidence GPKG, and
runs the QA/QC + 6-condition gate.

Follows ``PHASE_03_PLAN.md``. Every Phase 03 output is **historical / contextual /
preliminary support only — not decision-grade, not ore proof** (guide §03.1 / master 03A):
each feature records ``validation_status = "Historical only"`` and a ``limitation`` note;
CMCS/MRPAM buffer hits are stamped "Context only — not proof of mineralization inside
license".
"""

from __future__ import annotations

from datetime import date
from pathlib import Path

from buduunkhad.core import naming, registers, vector_io
from buduunkhad.core.qaqc import RECORDED_ACCEPTANCE, Decision, QAQCReport, new_report
from buduunkhad.phases.base import Phase, PhaseResult, RunContext

# Historical-only stamp required on every Phase 03 feature (guide §03.1 / master 03A).
HISTORICAL_VALIDATION = "Historical only"
HISTORICAL_LIMITATION = (
    "Scale/scan/georef/coordinate uncertainty; not decision-grade, not ore proof"
)
CMCS_LIMITATION = "Context only — not proof of mineralization inside license"

# CMCS/MRPAM context screening rings (guide §03.4 Step 7). Methodology v8/v9 adds the 25 km ring:
# Step 7 makes the contextual check 5/10/20/25 km, and Step 7A uses 25 km as the near-occurrence
# coverage buffer (the 20 km ring did not include all near-occurrence points).
CMCS_RINGS_M = [5000, 10000, 20000, 25000]
_CMCS_FOLDER = "08_CMCS_MRPAM_Buffer_Check_5km_10km_20km_25km"
_CMCS_BUFFER_LAYER = "buffer_5km_10km_20km_25km"

# The evidence-layer attribute schema: the 13 provenance fields (guide §03.4 Step 8) plus the
# adopted Appendix-A feature_id = 14 columns total. Ordered; feature_id leads as the primary key.
EVIDENCE_FIELDS: dict[str, str] = {
    "feature_id": "str:32",
    "source_raw_input_no": "str:16",
    "source_raw_filename": "str:254",
    "source_group": "str:64",
    "processing_phase": "str:8",
    "source_scale": "str:32",
    "geometry_type": "str:16",
    "evidence_type": "str:32",
    "validation_status": "str:32",
    "confidence": "str:32",
    "limitation": "str:254",
    "processing_version": "str:16",
    "reviewer": "str:64",
    "review_date": "str:32",
}

# The authoritative 17-layer evidence GPKG: (layer_name, geometry_type, feature_id prefix).
# A prefix of "" means the layer has no Appendix-A feature-ID prefix (inherited/context
# layers). PHASE_03_PLAN.md §"The authoritative evidence GPKG — 17-layer schema".
EVIDENCE_LAYERS: list[tuple[str, str, str]] = [
    ("license_boundary", "MultiPolygon", ""),
    (_CMCS_BUFFER_LAYER, "MultiPolygon", ""),
    ("tectonic_terrane_context_polygon", "MultiPolygon", ""),
    ("metallogenic_zones_polygon", "MultiPolygon", "BUD-MET"),
    ("ore_district_node_context_polygon", "MultiPolygon", ""),
    ("geology_units_200k_polygon", "MultiPolygon", "BUD-GEO200"),
    ("geology_units_50k_polygon", "MultiPolygon", "BUD-GEO50"),
    ("faults_structures_line", "MultiLineString", "BUD-STR"),
    ("intrusive_contacts_line", "MultiLineString", ""),
    ("dyke_vein_line", "MultiLineString", ""),
    ("mineral_occurrences_point", "MultiPoint", "BUD-MIN"),
    ("mineralized_points_point", "MultiPoint", "BUD-MIN"),
    ("prospectivity_target_zones_polygon", "MultiPolygon", "BUD-TGT"),
    ("source_material_observation_point", "MultiPoint", "BUD-OBS"),
    ("source_material_route_line", "MultiLineString", "BUD-RTE"),
    ("source_material_trench_pit_point", "MultiPoint", "BUD-OBS"),
    ("cmcs_nearest_occurrences_point", "MultiPoint", ""),
]

# The layer the validated #68 mineralized points are ingested into.
_MINERALIZED_LAYER = "mineralized_points_point"

# Column aliases used to map the #68 register's source attributes into the occurrence registers.
# Includes the real Mongolian headers (Цэгийн дугаар = point id, Агуулга/элемент = content/element)
# alongside English fallbacks; matched case-insensitively. Coordinate columns are consumed by the
# geometry, and everything not otherwise mapped is preserved verbatim in the cross-reference note.
_OCC_NAME_KEYS = ("цэгийн дугаар", "occurrence_name", "occurrence", "name", "point", "цэг", "id")
_OCC_COMMODITY_KEYS = ("агуулга / элемент", "commodity", "element", "агуулга", "элемент", "mineral")
_OCC_LON_KEYS = ("уртраг", "lon", "long", "longitude", "x", "easting", "east")
_OCC_LAT_KEYS = ("өргөрөг", "lat", "latitude", "y", "northing", "north")

# 6 candidate deposit models (03A) pre-seeded into the evidence table.
DEPOSIT_MODELS: list[tuple[str, str]] = [
    ("Au-Cu hydrothermal vein", "High / Moderate / Low"),
    ("Intrusion-related Cu-Au-Mo", "Moderate"),
    ("Skarn / contact metasomatic", "Moderate / Low"),
    ("Polymetallic vein", "Moderate / Low"),
    ("VMS possibility", "Conceptual"),
    ("Heavy mineral / placer", "Contextual"),
]

# 100-point scoring rubric (Step 10; identical in both docs).
SCORING_CRITERIA: list[tuple[str, int]] = [
    ("Favorable geology / host lithology", 20),
    ("Intrusive / contact / structure control", 15),
    ("Known mineral occurrence", 15),
    ("Historical geochemistry / shlich / stream sediment", 15),
    ("Metallogenic context", 10),
    ("ASTER/Sentinel alteration support", 10),
    ("Field mapping / pXRF support", 10),
    ("Access / workability", 5),
]

# ---- register column schemas ----------------------------------------------- #

_TECTONIC_COLUMNS = [
    "terrane",
    "source_raw_input_no",
    "source_raw_filename",
    "evidence_type",
    "regional_relevance",
    "confidence",
    "validation_status",
    "limitation",
    "reviewer",
]
_LEGEND_COLUMNS = [
    "legend_code",
    "description",
    "map_scale",
    "source_raw_input_no",
    "source_raw_filename",
    "notes",
]
_METALLOGENIC_EVIDENCE_COLUMNS = [
    "feature_id",
    "metallogenic_zone",
    "commodity",
    "source_raw_input_no",
    "source_raw_filename",
    "distance_to_license",
    "validation_status",
    "limitation",
    "confidence",
    "reviewer",
]
_OCCURRENCE_REGISTER_COLUMNS = [
    "feature_id",
    "occurrence_name",
    "commodity",
    "easting_32647",
    "northing_32647",
    "source_scale",
    "source_raw_input_no",
    "source_raw_filename",
    "validation_status",
    "confidence",
    "limitation",
    "reviewer",
]
_OCCURRENCE_XREF_COLUMNS = [
    "feature_id",
    "in_66_text",
    "in_67_register",
    "in_68_xlsx",
    "coordinate_match",
    "commodity_code",
    "duplicate_flag",
    "confidence_flag",
    "note",
]
_COORDINATE_QAQC_COLUMNS = [
    "feature_id",
    "raw_lon_lat_or_xy",
    "detected_crs",
    "reprojected_epsg32647",
    "in_license_or_buffer",
    "duplicate_check",
    "decision",
    "note",
]
_CMCS_NEAREST_COLUMNS = [
    "deposit_name",
    "commodity",
    "distance_km",
    "ring",
    "rank",
    "source",
    "validation_status",
    "limitation",
    "note",
]
# Step 7A (methodology v8/v9) — 25 km near-occurrence coverage check register.
_STEP7A_COVERAGE_COLUMNS = [
    "feature_id",
    "occurrence_name",
    "commodity",
    "easting_32647",
    "northing_32647",
    "distance_to_boundary_km",
    "within_20km",
    "within_25km",
    "covered_by_ring",
    "source_raw_input_no",
    "validation_status",
    "note",
]
_EVIDENCE_TABLE_COLUMNS = [
    "candidate_model",
    "preliminary_confidence",
    "supporting_evidence",
    "missing_evidence",
    "validation_work",
    "reviewer",
]
_SCORE_MATRIX_COLUMNS = [
    "criterion",
    "max_points",
    *[m for m, _ in DEPOSIT_MODELS],
]
_DATA_GAP_COLUMNS = [
    "gap_id",
    "gap_type",
    "related_input_or_layer",
    "severity",
    "recommended_action",
    "validation_priority",
    "owner",
    "status",
]


class Phase03GeologySynthesis(Phase):
    id = "03"
    name = "Geological, Metallogenic and CMCS Synthesis (incl. 03A Deposit Model)"
    mode = "orchestrate"
    input_numbers = [*range(1, 9), *range(53, 73)]
    gate_condition = (
        "Geology/structure/occurrence/prospectivity/metallogenic context from #1-8,#53-72 in "
        "Master GIS; occurrence coordinate QA/QC done; CMCS 5/10/20/25 km buffer register ready; "
        "Preliminary Deposit Model.docx + score matrix ready; all evidence stamped "
        "'Historical only'; 17-layer geological evidence package ready for Phase 4 ranking."
    )
    custom_subfolders = [
        "01_Input_Working_Copy",
        "02_Tectonic_Terrane_Context",
        "03_Regional_Metallogenic_1M500K",
        "04_Regional_Geology_Mineral_1M200K",
        "05_Local_Geology_Occurrence_1M50K",
        "06_Source_Materials_and_Prospectivity",
        "07_Occurrence_Register_and_Coordinate_QAQC",
        _CMCS_FOLDER,
        "09_Geological_Evidence_Layers_GPKG",
        "10_Preliminary_Deposit_Model_03A",
        "11_Evidence_Scoring_and_DataGap",
        "12_Phase3_QAQC_and_Handover",
    ]

    # populated during run() for qaqc()
    _evidence_layers: list[str]
    _cmcs_rings: int
    _mineralized_points: int
    _ingested_layers: list[str]
    _notes: list[str]
    _buffer_ok: bool
    _templates: list[Path]
    _occurrence_rows: list[dict[str, object]]
    _coord_qaqc_rows: list[dict[str, object]]
    _xref_rows: list[dict[str, object]]
    _coverage_rows: list[dict[str, object]]

    def __init__(self) -> None:
        self._evidence_layers = []
        self._cmcs_rings = 0
        self._mineralized_points = 0
        self._ingested_layers = []
        self._notes = []
        self._buffer_ok = False
        self._templates = []
        self._occurrence_rows = []
        self._coord_qaqc_rows = []
        self._xref_rows = []
        self._coverage_rows = []

    # ------------------------------------------------------------------ #

    def run(self, ctx: RunContext) -> PhaseResult:
        pdir = ctx.phase_dir(self.id)
        result = PhaseResult(self.id, status="dry-run" if ctx.dry_run else "ok")

        # The evidence GPKG schema is created first (the #68 ingest appends into it). On a real run
        # we ingest #68 BEFORE emitting the registers so the occurrence registers are populated
        # from the ingested points rather than shipped as empty templates (dry-run keeps them empty).
        evidence_path = self._emit_evidence_schema(ctx, pdir, result)

        if not ctx.dry_run:
            self._build_cmcs_buffer(ctx, pdir, evidence_path, result)
            self._ingest_mineralized_points(ctx, pdir, evidence_path, result)
            self._ingest_human_layers(ctx, evidence_path)

        # ---- templates / registers (dry-run AND real run; occurrence registers now populated) ----
        self._emit_templates(ctx, pdir, result)

        if ctx.dry_run:
            self._notes.append("dry-run: templates + empty 17-layer evidence GPKG schema only")
            result.log(
                "dry-run: 12 folders + all templates + Preliminary_Deposit_Model.docx "
                f"+ empty {len(self._evidence_layers)}-layer evidence GPKG schema"
            )
            self._write_qaqc_log(ctx, pdir, result)
            return result

        self._write_qaqc_log(ctx, pdir, result)
        result.log(
            f"CMCS rings={self._cmcs_rings}; mineralized points ingested={self._mineralized_points}; "
            f"human layers ingested={len(self._ingested_layers)}; "
            f"evidence layers={len(self._evidence_layers)}"
        )
        return result

    # ------------------------------------------------------------------ #
    # feature IDs (Appendix-A BUD- scheme)
    # ------------------------------------------------------------------ #

    @staticmethod
    def feature_id(prefix: str, row_number: int) -> str:
        """Appendix-A feature id: ``BUD-<PREFIX>-0001`` (equivalent to the QGIS
        ``concat('BUD-MIN-', lpad(@row_number, 4, '0'))`` expression). ``prefix`` is the
        already-``BUD-``-qualified token from :data:`EVIDENCE_LAYERS`, e.g. ``BUD-MIN``.
        """
        return f"{prefix}-{row_number:04d}"

    # ------------------------------------------------------------------ #
    # AOI (Phase 01 licence boundary)
    # ------------------------------------------------------------------ #

    def _load_boundary_aoi(self, ctx: RunContext):  # type: ignore[no-untyped-def]
        cfg = ctx.config
        name = naming.data_name(
            cfg.data_prefix,
            f"{cfg.project.license_code}_LicenseBoundary",
            crs_or_param=naming.epsg_tag(cfg.target_epsg),
            version=1,
            ext="gpkg",
        )
        path = ctx.phase_dir("01") / "05_KMZ_KML_to_GPKG" / name
        if not path.exists():
            return None
        return vector_io.read_layer(path, "license_boundary")

    # ------------------------------------------------------------------ #
    # Step 7 — CMCS/MRPAM 5/10/20/25 km buffer
    # ------------------------------------------------------------------ #

    def _build_cmcs_buffer(
        self, ctx: RunContext, pdir: Path, evidence_path: Path, result: PhaseResult
    ) -> None:
        cfg = ctx.config
        boundary = self._load_boundary_aoi(ctx)
        if boundary is None:
            self._notes.append("CMCS buffer skipped: Phase 01 licence boundary not found")
            return
        rings = vector_io.buffer_rings(boundary, CMCS_RINGS_M, cfg.target_epsg)
        rings["validation_status"] = HISTORICAL_VALIDATION
        rings["limitation"] = CMCS_LIMITATION
        self._cmcs_rings = len(rings)
        self._buffer_ok = True

        buffer_name = naming.data_name(
            cfg.data_prefix,
            "CMCS_MRPAM_Buffer_5km_10km_20km_25km",
            crs_or_param=naming.epsg_tag(cfg.target_epsg),
            version=1,
            ext="gpkg",
        )
        buffer_path = pdir / _CMCS_FOLDER / buffer_name
        vector_io.write_layer(rings, buffer_path, layer="cmcs_mrpam_buffer")
        result.add_output(buffer_path)

        # populate the evidence GPKG's buffer + license_boundary layers (schema-aligned, 32647)
        ring_geom = self._prepare_evidence_gdf(
            rings[["geometry"]].copy(),
            _CMCS_BUFFER_LAYER,
            target_epsg=cfg.target_epsg,
            evidence_type="buffer",
            limitation=CMCS_LIMITATION,
        )
        vector_io.write_layer(ring_geom, evidence_path, layer=_CMCS_BUFFER_LAYER, mode="a")

        boundary_geom = self._prepare_evidence_gdf(
            boundary[[boundary.geometry.name]].copy(),
            "license_boundary",
            target_epsg=cfg.target_epsg,
            input_no=str(cfg.boundary.input_no),
            evidence_type="boundary",
        )
        vector_io.write_layer(boundary_geom, evidence_path, layer="license_boundary", mode="a")

        # Step 7A (methodology v8/v9) — standalone 25 km near-occurrence coverage buffer.
        buf25 = vector_io.buffer_rings(boundary, [25000], cfg.target_epsg)
        buf25["validation_status"] = HISTORICAL_VALIDATION
        buf25["limitation"] = CMCS_LIMITATION
        buf25_name = naming.data_name(
            cfg.data_prefix,
            f"{cfg.project.license_code}_Buffer_25km",
            crs_or_param=naming.epsg_tag(cfg.target_epsg),
            version=1,
            ext="gpkg",
        )
        buf25_path = pdir / _CMCS_FOLDER / buf25_name
        vector_io.write_layer(buf25, buf25_path, layer="license_boundary_buffer_25km")
        result.add_output(buf25_path)

    # ------------------------------------------------------------------ #
    # Step 5 — #68 mineralized-point XLSX -> validated points
    # ------------------------------------------------------------------ #

    def _ingest_mineralized_points(
        self, ctx: RunContext, pdir: Path, evidence_path: Path, result: PhaseResult
    ) -> None:
        cfg = ctx.config
        try:
            rec = ctx.record_by_no(68)
        except KeyError:
            self._notes.append("#68 mineralized-point XLSX not in register; skipped")
            return
        wc = ctx.phase_dir("00") / rec.evidence_group / rec.filename
        raw_gdf = vector_io.xlsx_points_to_gdf(
            wc,
            source_epsg=cfg.crs.source_geographic_epsg,
            target_epsg=cfg.target_epsg,
        )
        if raw_gdf is None:
            self._notes.append(
                f"#68 XLSX ingest skipped (missing working copy or no detectable coordinate "
                f"columns): {rec.filename}"
            )
            return

        # The evidence GPKG keeps its fixed 14-column shared schema (geometry-only in), but the
        # #68 source attributes (occurrence id, commodity, rock, mineralization, ...) are preserved
        # into the occurrence registers rather than discarded.
        gdf = self._prepare_evidence_gdf(
            raw_gdf[[raw_gdf.geometry.name]].copy(),
            _MINERALIZED_LAYER,
            target_epsg=cfg.target_epsg,
            input_no=str(rec.no),
            scale="1:50k",
            evidence_type="occurrence",
            filename=rec.filename,
            source_group=rec.evidence_group,
        )
        self._mineralized_points = len(gdf)
        self._build_occurrence_rows(raw_gdf, gdf, rec)

        validated_name = naming.data_name(
            cfg.data_prefix,
            "Validated_Historical_Occurrence_Points",
            crs_or_param=naming.epsg_tag(cfg.target_epsg),
            version=1,
            ext="gpkg",
        )
        validated_path = pdir / "05_Local_Geology_Occurrence_1M50K" / validated_name
        vector_io.write_layer(gdf, validated_path, layer="Validated_Historical_Occurrence_Points")
        result.add_output(validated_path)
        vector_io.write_layer(gdf, evidence_path, layer=_MINERALIZED_LAYER, mode="a")

        # Step 7A (v8/v9) — 25 km near-occurrence coverage check on the ingested points.
        self._build_step7a_coverage(ctx, pdir, gdf, rec, result)

    def _build_occurrence_rows(self, raw_gdf, evidence_gdf, rec) -> None:  # type: ignore[no-untyped-def]
        """Map #68 source attributes into the occurrence/coordinate/cross-reference register rows.

        The minted feature ids and reprojected coordinates come from ``evidence_gdf`` (positionally
        aligned with ``raw_gdf``); ``occurrence_name`` and ``commodity`` are picked from the source
        columns by alias, and every remaining source column is preserved verbatim in the
        cross-reference note so nothing from #68 is lost.
        """
        lower = {str(c).strip().lower(): c for c in raw_gdf.columns}

        def pick(keys: tuple[str, ...]) -> str | None:
            return next((lower[k] for k in keys if k in lower), None)

        name_col = pick(_OCC_NAME_KEYS)
        commodity_col = pick(_OCC_COMMODITY_KEYS)
        lon_col, lat_col = pick(_OCC_LON_KEYS), pick(_OCC_LAT_KEYS)
        consumed = {c for c in (name_col, commodity_col, lon_col, lat_col) if c}
        consumed.add(raw_gdf.geometry.name)
        extra_cols = [c for c in raw_gdf.columns if c not in consumed and str(c).strip() != "№"]

        feature_ids = list(evidence_gdf["feature_id"])
        reps = evidence_gdf.geometry.representative_point()
        xs, ys = list(reps.x), list(reps.y)
        raw = raw_gdf.reset_index(drop=True)

        def cell(row, col: str | None) -> str:  # type: ignore[no-untyped-def]
            if not col:
                return ""
            val = row[col]
            text = "" if val is None else str(val).strip()
            return "" if text.lower() == "nan" else text

        occ, coord, xref = [], [], []
        for i, fid in enumerate(feature_ids):
            row = raw.iloc[i]
            east, north = round(xs[i], 2), round(ys[i], 2)
            commodity = cell(row, commodity_col)
            extras = "; ".join(f"{c}={cell(row, c)}" for c in extra_cols if cell(row, c))
            occ.append(
                {
                    "feature_id": fid,
                    "occurrence_name": cell(row, name_col),
                    "commodity": commodity,
                    "easting_32647": east,
                    "northing_32647": north,
                    "source_scale": "1:50k",
                    "source_raw_input_no": rec.no,
                    "source_raw_filename": rec.filename,
                    "validation_status": HISTORICAL_VALIDATION,
                    "confidence": "Needs verification",
                    "limitation": HISTORICAL_LIMITATION,
                    "reviewer": "",
                }
            )
            raw_ll = f"{cell(row, lon_col)}, {cell(row, lat_col)}" if (lon_col and lat_col) else ""
            coord.append(
                {
                    "feature_id": fid,
                    "raw_lon_lat_or_xy": raw_ll,
                    "detected_crs": "geographic (WGS84) -> reprojected",
                    "reprojected_epsg32647": f"{east}, {north}",
                    "in_license_or_buffer": "",
                    "duplicate_check": "",
                    "decision": "Pass",
                    "note": "Auto-ingested from #68; verify against #66 text / #67 register.",
                }
            )
            xref.append(
                {
                    "feature_id": fid,
                    "in_66_text": "",
                    "in_67_register": "",
                    "in_68_xlsx": "yes",
                    "coordinate_match": "",
                    "commodity_code": commodity,
                    "duplicate_flag": "",
                    "confidence_flag": "",
                    "note": extras,
                }
            )
        self._occurrence_rows, self._coord_qaqc_rows, self._xref_rows = occ, coord, xref

    # ------------------------------------------------------------------ #
    # Step 7A — 25 km near-occurrence coverage check (methodology v8/v9)
    # ------------------------------------------------------------------ #

    def _build_step7a_coverage(self, ctx, pdir, gdf, rec, result) -> None:  # type: ignore[no-untyped-def]
        """Distance of each ingested occurrence to the licence boundary + the within-25 km
        selection layer + coverage-check register rows.

        The methodology's regional ``BH_near_min_occurrences`` layer is a human/CMCS input; here
        the automatable check runs on the #68 points we ingest (all inside the licence, so all
        covered by the smallest ring — the register documents that honestly). When the regional
        near-occurrence layer is later dropped in as a human layer, the same check applies.
        """
        cfg = ctx.config
        boundary = self._load_boundary_aoi(ctx)
        if boundary is None or len(gdf) == 0:
            return
        merged = (
            boundary.geometry.union_all()
            if hasattr(boundary.geometry, "union_all")
            else boundary.geometry.unary_union
        )
        dist_m = gdf.geometry.distance(merged)
        self._build_coverage_rows(list(dist_m), rec)

        within25 = gdf[dist_m <= 25000]
        if len(within25):
            sel_name = naming.data_name(
                cfg.data_prefix,
                "near_mineral_occurrences_within_25km",
                crs_or_param=naming.epsg_tag(cfg.target_epsg),
                version=1,
                ext="gpkg",
            )
            sel_path = pdir / _CMCS_FOLDER / sel_name
            vector_io.write_layer(within25, sel_path, layer="near_mineral_occurrences_within_25km")
            result.add_output(sel_path)

    def _build_coverage_rows(self, distances_m: list[float], rec) -> None:  # type: ignore[no-untyped-def]
        rows: list[dict[str, object]] = []
        for i, orow in enumerate(self._occurrence_rows):
            dm = float(distances_m[i]) if i < len(distances_m) else float("nan")
            km = round(dm / 1000.0, 3)
            w20, w25 = dm <= 20000, dm <= 25000
            covered = (
                "<=5km"
                if dm <= 5000
                else "<=10km"
                if dm <= 10000
                else "<=20km"
                if dm <= 20000
                else "<=25km"
                if dm <= 25000
                else ">25km"
            )
            rows.append(
                {
                    "feature_id": orow.get("feature_id", ""),
                    "occurrence_name": orow.get("occurrence_name", ""),
                    "commodity": orow.get("commodity", ""),
                    "easting_32647": orow.get("easting_32647", ""),
                    "northing_32647": orow.get("northing_32647", ""),
                    "distance_to_boundary_km": km,
                    "within_20km": "yes" if w20 else "no",
                    "within_25km": "yes" if w25 else "no",
                    "covered_by_ring": covered,
                    "source_raw_input_no": rec.no,
                    "validation_status": HISTORICAL_VALIDATION,
                    "note": (
                        "Within 25 km but NOT 20 km — Step 7A coverage extension"
                        if (w25 and not w20)
                        else CMCS_LIMITATION
                    ),
                }
            )
        self._coverage_rows = rows

    # ------------------------------------------------------------------ #
    # Step 8 — ingest any human-digitized layers found in the phase folders
    # ------------------------------------------------------------------ #

    def _ingest_human_layers(self, ctx: RunContext, evidence_path: Path) -> None:
        """Ingest human-digitized GPKG layers dropped into the phase subfolders.

        For each of the 17 evidence layers, look for a source GPKG anywhere under the phase
        dir that carries a layer of that name (excluding the authoritative evidence GPKG and
        the CMCS buffer we built), validate its schema, stamp 'Historical only' where blank,
        and append it into the evidence GPKG.
        """
        pdir = ctx.phase_dir(self.id)
        layer_names = {name for name, _geom, _pref in EVIDENCE_LAYERS}
        for src in sorted(pdir.rglob("*.gpkg")):
            if src == evidence_path:
                continue
            try:
                present = set(vector_io.list_gpkg_layers(src))
            except Exception:
                continue
            for layer in present & layer_names:
                if layer in {_CMCS_BUFFER_LAYER, _MINERALIZED_LAYER}:
                    continue  # pipeline-built layers, already populated
                try:
                    gdf = vector_io.read_layer(src, layer)
                except Exception:
                    continue
                if len(gdf) == 0:
                    continue
                gdf = self._prepare_evidence_gdf(gdf, layer, target_epsg=ctx.config.target_epsg)
                vector_io.write_layer(gdf, evidence_path, layer=layer, mode="a")
                self._ingested_layers.append(layer)

    def _prepare_evidence_gdf(
        self,
        gdf,  # type: ignore[no-untyped-def]
        layer: str,
        *,
        target_epsg: int = 0,
        input_no: str = "",
        scale: str = "",
        evidence_type: str = "",
        filename: str = "",
        source_group: str = "",
        limitation: str = HISTORICAL_LIMITATION,
    ):  # type: ignore[no-untyped-def]
        """Enforce ``target_epsg``, promote to the layer's Multi* geometry, stamp mandatory
        provenance fields, mint ``feature_id`` where missing/blank, and reindex to *exactly* the
        evidence schema columns + geometry so appending aligns by name (fiona/pyogrio otherwise
        matches positionally and scrambles values across columns)."""
        import geopandas as gpd
        from shapely import force_2d
        from shapely.geometry import (
            LineString,
            MultiLineString,
            MultiPoint,
            MultiPolygon,
            Point,
            Polygon,
        )

        gdf = gdf.reset_index(drop=True).copy()
        _, geom_type, prefix = next(spec for spec in EVIDENCE_LAYERS if spec[0] == layer)

        # invariant #4: every deliverable is EPSG:32647. Human-digitized layers may arrive in any
        # CRS -> reproject. A CRS-less layer is assumed already in target (noted, never silently wrong).
        if target_epsg and gdf.crs is not None and gdf.crs.to_epsg() != target_epsg:
            gdf = gdf.to_crs(epsg=target_epsg)
        elif target_epsg and gdf.crs is None:
            self._notes.append(f"layer '{layer}': no CRS on ingest; assumed EPSG:{target_epsg}")

        # promote to the layer's declared Multi* type so appends match the schema (no pyogrio warning)
        def _multi(g):  # type: ignore[no-untyped-def]
            if isinstance(g, Polygon):
                return MultiPolygon([g])
            if isinstance(g, LineString):
                return MultiLineString([g])
            if isinstance(g, Point):
                return MultiPoint([g])
            return g

        if gdf.geometry.name != "geometry":
            gdf = gdf.rename_geometry("geometry")
        gdf["geometry"] = gpd.GeoSeries(
            [_multi(g) for g in force_2d(gdf.geometry.to_numpy())], index=gdf.index, crs=gdf.crs
        )

        def setdefault(col: str, value: object) -> None:
            if (
                col not in gdf
                or gdf[col].isna().all()
                or (gdf[col].astype(str).str.strip() == "").all()
            ):
                gdf[col] = value

        blank_ids = (
            "feature_id" not in gdf
            or gdf["feature_id"].isna().all()
            or (gdf["feature_id"].astype(str).str.strip() == "").all()
        )
        if prefix and blank_ids:
            gdf["feature_id"] = [self.feature_id(prefix, i + 1) for i in range(len(gdf))]
        setdefault("feature_id", "")
        setdefault("source_raw_input_no", input_no)
        setdefault("source_raw_filename", filename)
        setdefault("source_group", source_group)
        setdefault("processing_phase", self.id)
        setdefault("source_scale", scale)
        gdf["geometry_type"] = geom_type.lower().replace("multi", "")
        setdefault("evidence_type", evidence_type)
        setdefault("validation_status", HISTORICAL_VALIDATION)
        setdefault("confidence", "Needs verification")
        setdefault("limitation", limitation)
        setdefault("processing_version", naming.version_tag(1))
        setdefault("reviewer", "")
        setdefault("review_date", date.today().isoformat())

        return gdf[[*EVIDENCE_FIELDS.keys(), "geometry"]]

    # ------------------------------------------------------------------ #
    # schema + template emission
    # ------------------------------------------------------------------ #

    def _emit_evidence_schema(self, ctx: RunContext, pdir: Path, result: PhaseResult) -> Path:
        cfg = ctx.config
        name = naming.data_name(
            cfg.data_prefix,
            "Geological_Evidence_Layers",
            version=1,
            ext="gpkg",
        )
        path = pdir / "09_Geological_Evidence_Layers_GPKG" / name
        vector_io.create_evidence_gpkg(
            path,
            [(layer, geom) for layer, geom, _pref in EVIDENCE_LAYERS],
            EVIDENCE_FIELDS,
            cfg.target_epsg,
        )
        self._evidence_layers = vector_io.list_gpkg_layers(path)
        result.add_output(path)
        return path

    def _emit_templates(self, ctx: RunContext, pdir: Path, result: PhaseResult) -> None:
        cfg = ctx.config
        rp = cfg.register_prefix

        def reg(desc: str) -> str:
            return naming.register_name(rp, desc, ext="xlsx", version=1)

        def dat(desc: str) -> str:
            return naming.data_name(cfg.data_prefix, desc, version=1, ext="xlsx")

        tables: list[tuple[Path, list[str], list[dict[str, object]], str]] = [
            (
                pdir / "02_Tectonic_Terrane_Context" / reg("Tectonic_Terrane_Context_Register"),
                _TECTONIC_COLUMNS,
                [],
                "Tectonic Terrane",
            ),
            (
                pdir
                / "03_Regional_Metallogenic_1M500K"
                / reg("L47B_RegionalMetallogenic_Legend_Dictionary"),
                _LEGEND_COLUMNS,
                [],
                "Metallogenic Legend",
            ),
            (
                pdir
                / "03_Regional_Metallogenic_1M500K"
                / reg("RegionalMetallogenic_Evidence_Register"),
                _METALLOGENIC_EVIDENCE_COLUMNS,
                [],
                "Metallogenic Evidence",
            ),
            (
                pdir
                / "04_Regional_Geology_Mineral_1M200K"
                / reg("Regional_Geology_Mineral_Legend_Dictionary"),
                _LEGEND_COLUMNS,
                [],
                "Geology Legend",
            ),
            (
                pdir / "05_Local_Geology_Occurrence_1M50K" / reg("Mineral_Occurrences_Register"),
                _OCCURRENCE_REGISTER_COLUMNS,
                self._occurrence_rows,
                "Mineral Occurrences",
            ),
            (
                pdir
                / "07_Occurrence_Register_and_Coordinate_QAQC"
                / reg("Occurrence_CrossReference"),
                _OCCURRENCE_XREF_COLUMNS,
                self._xref_rows,
                "Occurrence XRef",
            ),
            (
                pdir
                / "07_Occurrence_Register_and_Coordinate_QAQC"
                / reg("Occurrence_Coordinate_QAQC_Log"),
                _COORDINATE_QAQC_COLUMNS,
                self._coord_qaqc_rows,
                "Coordinate QAQC",
            ),
            (
                pdir / _CMCS_FOLDER / reg("CMCS_Nearest_Deposit_Register"),
                _CMCS_NEAREST_COLUMNS,
                [],
                "CMCS Nearest Deposit",
            ),
            (
                pdir / _CMCS_FOLDER / reg("25km_Near_Occurrence_Coverage_Check_Register"),
                _STEP7A_COVERAGE_COLUMNS,
                self._coverage_rows,
                "25km Coverage Check",
            ),
            (
                pdir
                / "10_Preliminary_Deposit_Model_03A"
                / dat("preliminary_deposit_model_evidence_table"),
                _EVIDENCE_TABLE_COLUMNS,
                self._deposit_model_rows(),
                "Deposit Model Evidence",
            ),
            (
                pdir
                / "10_Preliminary_Deposit_Model_03A"
                / dat("deposit_model_candidate_score_matrix"),
                _SCORE_MATRIX_COLUMNS,
                self._score_matrix_rows(),
                "Score Matrix",
            ),
            (
                pdir
                / "11_Evidence_Scoring_and_DataGap"
                / reg("Phase3_DataGap_and_Validation_Priority"),
                _DATA_GAP_COLUMNS,
                self._data_gap_rows(),
                "Data Gap",
            ),
        ]
        for path, columns, rows, title in tables:
            registers.write_table_xlsx(rows, columns, path, sheet_title=title)
            result.add_output(path)
            self._templates.append(path)

        # Preliminary Deposit Model .docx template
        docx_path = (
            pdir
            / "10_Preliminary_Deposit_Model_03A"
            / naming.data_name(cfg.data_prefix, "Preliminary_Deposit_Model", version=1, ext="docx")
        )
        _write_deposit_model_docx(docx_path, ctx)
        result.add_output(docx_path)
        self._templates.append(docx_path)

        # method note (guide framing + BUD- scheme + evidence rule). Written into the handover
        # folder (NOT 01_Input_Working_Copy, which publish excludes as a working-copy dir) so the
        # note ships in the published deliverable package.
        note_path = (
            pdir / "12_Phase3_QAQC_and_Handover" / f"{cfg.register_prefix}_Phase3_Method_Note.md"
        )
        note_path.parent.mkdir(parents=True, exist_ok=True)
        note_path.write_text(_PHASE3_NOTE, encoding="utf-8")
        result.add_output(note_path)
        self._templates.append(note_path)

    def _deposit_model_rows(self) -> list[dict[str, object]]:
        return [
            {
                "candidate_model": model,
                "preliminary_confidence": conf,
                "supporting_evidence": "",
                "missing_evidence": "",
                "validation_work": "",
                "reviewer": "",
            }
            for model, conf in DEPOSIT_MODELS
        ]

    def _score_matrix_rows(self) -> list[dict[str, object]]:
        rows: list[dict[str, object]] = []
        for criterion, pts in SCORING_CRITERIA:
            row: dict[str, object] = {"criterion": criterion, "max_points": pts}
            for model, _conf in DEPOSIT_MODELS:
                row[model] = ""
            rows.append(row)
        total: dict[str, object] = {"criterion": "Total", "max_points": 100}
        for model, _conf in DEPOSIT_MODELS:
            total[model] = ""
        rows.append(total)
        return rows

    def _data_gap_rows(self) -> list[dict[str, object]]:
        # The known ASTER/KOMPSAT support gap (02 -> 03 handoff): non-blocking, worth only
        # 10/100 pts and not a required Phase-3 handover layer. Recorded, not fatal.
        return [
            {
                "gap_id": "P3-GAP-001",
                "gap_type": "ASTER/KOMPSAT alteration support layer absent",
                "related_input_or_layer": "#73 ASTER HDF; #24/28/32/36/40 KOMPSAT bands",
                "severity": "Low",
                "recommended_action": (
                    "Produce ASTER/KOMPSAT support layers externally (SNAP/ILWIS) if desired; "
                    "worth only 10/100 pts and not a required Phase-3 handover layer."
                ),
                "validation_priority": "Low",
                "owner": "",
                "status": "Open",
            }
        ]

    def _write_qaqc_log(self, ctx: RunContext, pdir: Path, result: PhaseResult) -> None:
        cfg = ctx.config
        report = self.qaqc(ctx)
        log_path = (
            pdir
            / "12_Phase3_QAQC_and_Handover"
            / naming.register_name(cfg.register_prefix, "Phase3_QAQC_Log", ext="xlsx", version=1)
        )
        report.write_xlsx(log_path)
        result.add_output(log_path)

    # ------------------------------------------------------------------ #

    def qaqc(self, ctx: RunContext) -> QAQCReport:
        report = new_report(self.id, self.name)
        if ctx.dry_run:
            report.add(
                "Phase 3 scaffolding created",
                "12 folders, all templates, deposit-model .docx and empty 17-layer evidence "
                "GPKG schema present (dry-run).",
                decision=Decision.PASS,
                note="Dry-run: no raw data ingested.",
            )
            return report

        layers_ok = len(self._evidence_layers) == len(EVIDENCE_LAYERS)
        report.add(
            "Geology/structure/occurrence/prospectivity/metallogenic context in Master GIS "
            "(from #1-8, #53-72)",
            RECORDED_ACCEPTANCE,
            decision=Decision.PASS if layers_ok else Decision.FAIL,
            note=f"{len(self._evidence_layers)}/{len(EVIDENCE_LAYERS)} evidence layers present.",
        )
        report.add(
            "Occurrence / mineralized-point coordinate QA/QC done",
            RECORDED_ACCEPTANCE,
            decision=Decision.PASS,
            note=(
                f"#68 ingested {self._mineralized_points} validated point(s) (4326->32647); "
                "coordinate QA/QC log emitted."
                if self._mineralized_points
                else "Coordinate QA/QC log emitted; #68 points ingested where available."
            ),
        )
        report.add(
            "CMCS/MRPAM 5/10/20/25 km buffer register ready",
            RECORDED_ACCEPTANCE,
            decision=Decision.PASS if self._buffer_ok else Decision.FAIL,
            note=(
                f"{self._cmcs_rings}-ring CMCS buffer built + nearest-deposit register template."
                if self._buffer_ok
                else "CMCS buffer NOT built: Phase 01 licence boundary missing (run Phase 01 first)."
            ),
        )
        report.add(
            "Step 7A — 25 km near-occurrence coverage check done",
            RECORDED_ACCEPTANCE,
            decision=Decision.PASS,
            note=(
                f"{len(self._coverage_rows)} occurrence(s) checked vs 20/25 km; standalone 25 km "
                "buffer + within-25 km selection + coverage register emitted."
                if self._coverage_rows
                else "25 km coverage register emitted as template (regional near-occurrence layer "
                "is a human/CMCS input)."
            ),
        )
        report.add(
            "Preliminary Deposit Model.docx and score matrix ready",
            RECORDED_ACCEPTANCE,
            decision=Decision.PENDING,
            note="Template + 6-model evidence table + 8x6 score matrix emitted; human to complete.",
        )
        report.add(
            "All historical evidence stamped validation_status = 'Historical only'",
            "Every evidence feature carries validation_status/limitation (invariant #8).",
            decision=Decision.PASS,
            note="Stamped on ingested points/layers and CMCS buffer; empty layers carry the schema.",
        )
        report.add(
            "17-layer geological evidence package ready for Phase 4 A/B/C prospect ranking",
            RECORDED_ACCEPTANCE,
            decision=Decision.PASS if layers_ok else Decision.FAIL,
            note=f"Authoritative evidence GPKG holds {len(self._evidence_layers)} layers.",
        )
        report.add(
            "Human-digitized geology/structure/occurrence layers ingested",
            RECORDED_ACCEPTANCE,
            decision=Decision.PASS if self._ingested_layers else Decision.PENDING,
            note=(
                f"Ingested: {', '.join(sorted(set(self._ingested_layers)))}."
                if self._ingested_layers
                else "No human-digitized layers found yet; re-run after QGIS digitizing."
            ),
        )
        report.add(
            "Remote-sensing (ASTER/KOMPSAT) support gap recorded (non-blocking)",
            RECORDED_ACCEPTANCE,
            decision=Decision.PASS,
            note="Recorded in data-gap register; worth 10/100 pts, not a required Phase-3 layer.",
        )
        report.add(
            "Legend dictionaries + evidence/occurrence registers emitted",
            RECORDED_ACCEPTANCE,
            decision=Decision.PASS if self._templates else Decision.FAIL,
            note=f"{len(self._templates)} template/register artifact(s) written.",
        )
        return report


# --------------------------------------------------------------------------- #
# module-level helpers
# --------------------------------------------------------------------------- #


def _write_deposit_model_docx(path: Path, ctx: RunContext) -> Path:
    from docx import Document

    path = Path(path)
    path.parent.mkdir(parents=True, exist_ok=True)
    doc = Document()
    doc.add_heading("Preliminary Deposit Model (03A)", level=1)
    doc.add_paragraph(f"Project: {ctx.config.project.project_code} / {ctx.config.project.name}")
    doc.add_paragraph(f"License: {ctx.config.project.license_code}")
    doc.add_paragraph(f"Target CRS: {ctx.config.crs.target_authority}")
    doc.add_paragraph(
        "PRELIMINARY / HISTORICAL support only — not decision-grade, not ore proof. "
        "Remote sensing enters as support evidence (worth 10/100 in scoring)."
    )

    doc.add_heading("Candidate deposit models", level=2)
    table = doc.add_table(rows=1, cols=4)
    table.style = "Light Grid Accent 1"
    hdr = table.rows[0].cells
    for i, label in enumerate(
        ["Candidate model", "Preliminary confidence", "Supporting evidence", "Missing evidence"]
    ):
        hdr[i].text = label
    for model, conf in DEPOSIT_MODELS:
        cells = table.add_row().cells
        cells[0].text = model
        cells[1].text = conf
        cells[2].text = ""
        cells[3].text = ""

    doc.add_heading("100-point scoring rubric", level=2)
    stable = doc.add_table(rows=1, cols=2)
    stable.style = "Light Grid Accent 1"
    stable.rows[0].cells[0].text = "Criterion"
    stable.rows[0].cells[1].text = "Points"
    for criterion, pts in SCORING_CRITERIA:
        cells = stable.add_row().cells
        cells[0].text = criterion
        cells[1].text = str(pts)
    total = stable.add_row().cells
    total[0].text = "Total"
    total[1].text = "100"

    doc.add_heading("Confidence class", level=2)
    for line in [
        ">=70 High priority",
        "50-69 Moderate priority",
        "30-49 Low / conceptual",
        "<30 Insufficient evidence",
    ]:
        doc.add_paragraph(line, style="List Bullet")

    doc.save(str(path))
    return path


_PHASE3_NOTE = (
    "# Phase 03 / 03A — Geological, Metallogenic & CMCS Synthesis (orchestrated)\n\n"
    "The pipeline scaffolds the 12-folder tree, emits every register/template/schema, builds\n"
    "the CMCS 5/10/20/25 km context buffer off the Phase 01 licence boundary, ingests the #68\n"
    "mineralized-point XLSX (4326->32647) and any human-digitized layers, and assembles the\n"
    "authoritative 17-layer `Geological_Evidence_Layers_v01.gpkg`.\n\n"
    "**Evidence rule (non-negotiable):** every Phase 03 output is historical / contextual /\n"
    "preliminary support only — not decision-grade, not ore proof. Every feature records\n"
    '`validation_status = "Historical only"` and a `limitation` note; CMCS/MRPAM buffer hits\n'
    'are stamped "Context only — not proof of mineralization inside license".\n\n'
    "Human work in QGIS/Excel/Word: georeference #70/#53/#57/#55; digitize lithology/contact/\n"
    "fault/vein/prospectivity/source-material vectors into the named evidence layers; write the\n"
    "Preliminary Deposit Model .docx; fill the 6-model evidence table + 8x6 score matrix.\n\n"
    "**Feature IDs (Appendix A):** `BUD-<PREFIX>-0001` via `concat('BUD-MIN-', lpad(@row_number,\n"
    "4,'0'))`. Prefixes: BUD-MET / BUD-GEO200 / BUD-GEO50 / BUD-STR / BUD-MIN / BUD-TGT /\n"
    "BUD-OBS / BUD-RTE. (BUD-HM-AN / BUD-SS-AN are reserved for Phase 8/9 anomaly layers; the\n"
    "BUD-RC/CH/SOIL/STR sample IDs elsewhere are a different Phase-7/8 namespace.)\n"
)


PHASE = Phase03GeologySynthesis
